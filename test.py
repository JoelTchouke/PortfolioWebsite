from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement

# Function to create a well-formatted paragraph
def add_bullet_paragraph(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)

# Create a new document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

# Add header
doc.add_paragraph("Joel Tchouke", style="Heading 1").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Software Engineering Internship", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("yvesjoel.tchouke@mnsu.edu | linkedin.com/in/joel-tchouke-197390280/ | github.com/jolvie | joeltchouke.com", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Mankato, MN | (+1) 929-339-7034", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add sections
doc.add_paragraph("SUMMARY OF QUALIFICATIONS", style="Heading 2")
add_bullet_paragraph(doc, "Proficient in full-stack development with expertise in React, Django, Firebase, and C++, solving real-world problems efficiently.")
add_bullet_paragraph(doc, "Strong problem-solving skills demonstrated through cross-platform applications and AI-driven feature implementations.")
add_bullet_paragraph(doc, "Experienced in leading teams, collaborating with clients, and mentoring students in software engineering and cybersecurity.")

doc.add_paragraph("EDUCATION", style="Heading 2")
doc.add_paragraph("Minnesota State University, Mankato", style="Normal").bold = True
doc.add_paragraph("Computer Engineering | Honors Student | Dean’s List | Physics minor | 2026\nGPA: 3.7/4.0", style="Normal")

doc.add_paragraph("SKILLS / COURSES", style="Heading 2")
skills = [
    "C++, C#, Java, Python, Go, JavaScript, React.js, Node.js, Django, MongoDB, MySQL, SQLAlchemy, Machine Learning, TensorFlow, Pytorch, Panda, Git, HTML/CSS, TypeScript,Postgres, AWS SQS, Firebase, .NET, REST API, MatLab, Splunk",
    "Object Oriented Programming, Algorithms, Relational Database Design & SQL, Computer Systems, Linear Algebra, Discrete Mathematics, Applied Probability and Statistics, Computer Architecture",
    ".Net, WinRT, C++, C#, CICD"
]
for skill in skills:
    add_bullet_paragraph(doc, skill)

# Save the document
file_path = "/mnt/data/Joel_Tchouke_Resume.docx"
doc.save(file_path)
file_path
